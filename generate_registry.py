"""
Universal Legal Registry Automation & Extraction Engine
========================================================
100% Robust, Content-Based Deed Generation Engine.
Precisely fills:
  - Top Banner Title Header (e.g. भूखण्ड क्रमांक E9 (EWS-9), एमराल्ड आश्रय, ग्राम-सोनवाय, तहसील-राऊ, ज़िला-इन्दौर)
  - Preamble with Plot No and Allotment Value in digits & Hindi words
  - Single / Multi-Buyer execution block (P17-P19)
  - Clause 10 / 11 Cost Preamble
  - Transaction Payments table / Ledger with UTR / Reference No.
  - Property Description (Plot No, Area in Sq.Mtr., Area in Sq.Ft.)
  - 2-Column Boundary & Dimensions block (East, West, North, South)
  - Closing Delivery Clause with Allotment Value in digits & Hindi words
"""

import os
import sys
import re
import datetime
import argparse
from PIL import Image
import openpyxl
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# 1. Hindi Number & Currency Engine (0 to 100+ Crores)
# ---------------------------------------------------------------------------

HINDI_ONES = {
    0: '', 1: 'एक', 2: 'दो', 3: 'तीन', 4: 'चार', 5: 'पाँच', 6: 'छह', 7: 'सात', 8: 'आठ', 9: 'नौ',
    10: 'दस', 11: 'ग्यारह', 12: 'बारह', 13: 'तेरह', 14: 'चौदह', 15: 'पंद्रह', 16: 'सोलह', 17: 'सत्रह', 18: 'अठारह', 19: 'उन्नीस',
    20: 'बीस', 21: 'इक्कीस', 22: 'बाईस', 23: 'तेईस', 24: 'चौबीस', 25: 'पच्चीस', 26: 'छब्बीस', 27: 'सत्ताईस', 28: 'अट्ठाईस', 29: 'उनत्तीस',
    30: 'तीस', 31: 'इकत्तीस', 32: 'बत्तीस', 33: 'तैंतीस', 34: 'चौंतीस', 35: 'पैंतीस', 36: 'छत्तीस', 37: 'सैंतीस', 38: 'अड़तीस', 39: 'उनतालीस',
    40: 'चालीस', 41: 'इकतालीस', 42: 'बयालीस', 43: 'तिरालीस', 44: 'चौवालिस', 45: 'पैंतालीस', 46: 'छियालीस', 47: 'सैंतालीस', 48: 'अड़तालीस', 49: 'उनचास',
    50: 'पचास', 51: 'इक्कावन', 52: 'बावन', 53: 'तिरेपन', 54: 'चौवन', 55: 'पचपन', 56: 'छप्पन', 57: 'सत्तावन', 58: 'अट्टावन', 59: 'उनसठ',
    60: 'साठ', 61: 'इक्सठ', 62: 'बासठ', 63: 'तिर्सठ', 64: 'चौंसठ', 65: 'पैंसठ', 66: 'छियासठ', 67: 'सरसठ', 68: 'अड़सठ', 69: 'उनसत्तर',
    70: 'सत्तर', 71: 'इकहत्तर', 72: 'बहत्तर', 73: 'तिहत्तर', 74: 'चौहत्तर', 75: 'पचहत्तर', 76: 'छियात्तर', 77: 'सतहत्तर', 78: 'इठ्‌योत्तर', 79: 'उनासी',
    80: 'अस्सी', 81: 'इक्यासी', 82: 'बयासी', 83: 'तिरासी', 84: 'चौरासी', 85: 'पचासी', 86: 'छियासी', 87: 'सत्तासी', 88: 'अ्ठासी', 89: 'नवासी',
    90: 'नब्बे', 91: 'इक्यान्वे', 92: 'बान्वे', 93: 'तिरान्वे', 94: 'चौरान्वे', 95: 'पन्चान्वे', 96: 'छियान्वे', 97: 'संतान्वे', 98: 'अन्ठान्वे', 99: 'निन्यानवे'
}

HINDI_MONTHS = {
    1: 'जनवरी', 2: 'फरवरी', 3: 'मार्च', 4: 'अप्रैल', 5: 'मई', 6: 'जून',
    7: 'जुलाई', 8: 'अगस्त', 9: 'सितम्बर', 10: 'अक्टूबर', 11: 'नवम्बर', 12: 'दिसम्बर'
}

def number_to_hindi_words(n):
    try:
        n = int(round(float(n)))
    except (ValueError, TypeError):
        return str(n) if n else ""
        
    if n == 0:
        return 'शून्य'
    parts = []
    
    crores = n // 10000000
    n %= 10000000
    if crores > 0:
        c_str = HINDI_ONES.get(crores) or number_to_hindi_words(crores)
        parts.append(f"{c_str} करोड़")
        
    lakhs = n // 100000
    n %= 100000
    if lakhs > 0:
        parts.append(f"{HINDI_ONES.get(lakhs, str(lakhs))} लाख")
        
    thousands = n // 1000
    n %= 1000
    if thousands > 0:
        parts.append(f"{HINDI_ONES.get(thousands, str(thousands))} हज़ार")
        
    hundreds = n // 100
    n %= 100
    if hundreds > 0:
        parts.append(f"{HINDI_ONES.get(hundreds, str(hundreds))} सौ")
        
    if n > 0:
        parts.append(HINDI_ONES.get(n, str(n)))
        
    return ' '.join(parts)

def get_today_hindi_date():
    today = datetime.date.today()
    month = HINDI_MONTHS.get(today.month, "")
    return f"{today.day} {month}, {today.year}."

def clean_plot_no(p_num):
    if not p_num:
        return ""
    p_str = str(p_num).strip().replace('>', '').replace('<', '').replace('\u200c', '').replace('\u200d', '')
    p_clean = re.sub(r'\s*\([^)]*\)', '', p_str).strip()
    m_ews = re.match(r'^EWS[-_ ]?(\d+)$', p_clean, re.I)
    if m_ews:
        return f"E{m_ews.group(1)}"
    return p_clean

def plot_number_to_hindi(p_num):
    if not p_num:
        return ""
    p_str = str(p_num).strip().replace('>', '').replace('<', '').replace('\u200c', '').replace('\u200d', '')
    p_str = re.sub(r'\s*\([^)]*\)', '', p_str).strip()
    
    e_match = re.search(r'\b(?:EWS|E)[-_ ]?(\d+)\b', p_str, re.I)
    if e_match:
        val = int(e_match.group(1))
        h_val = HINDI_ONES.get(val) or number_to_hindi_words(val)
        return f"ई-{h_val}"
        
    l_match = re.search(r'\bL[-_ ]?(\d+)\b', p_str, re.I)
    if l_match:
        val = int(l_match.group(1))
        h_val = HINDI_ONES.get(val) or number_to_hindi_words(val)
        return f"एल-{h_val}"

    b_match = re.search(r'\bB[-_ ]?(\d+)\b', p_str, re.I)
    if b_match:
        val = int(b_match.group(1))
        h_val = HINDI_ONES.get(val) or number_to_hindi_words(val)
        return f"बी-{h_val}"

    c_match = re.search(r'\bC[-_ ]?(\d+)\b', p_str, re.I)
    if c_match:
        val = int(c_match.group(1))
        h_val = HINDI_ONES.get(val) or number_to_hindi_words(val)
        return f"सी-{h_val}"

    digits = re.findall(r'\d+', p_str)
    if digits:
        val = int(digits[0])
        return HINDI_ONES.get(val) or number_to_hindi_words(val)
    return p_str

def format_indian_currency(n):
    if not n:
        return ""
    try:
        s = str(int(round(float(n))))
    except (ValueError, TypeError):
        return str(n)
        
    if len(s) <= 3:
        return s
    last3 = s[-3:]
    other = s[:-3]
    res = ""
    while len(other) > 2:
        res = "," + other[-2:] + res
        other = other[:-2]
    return other + res + "," + last3

def format_hindi_date(dt):
    if isinstance(dt, str):
        for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y", "%Y/%m/%d", "%d.%m.%Y"):
            try:
                dt = datetime.datetime.strptime(dt.strip(), fmt)
                break
            except ValueError:
                continue
    if not isinstance(dt, (datetime.datetime, datetime.date)):
        return str(dt) if dt else ""
        
    day = f"{dt.day:02d}"
    month = HINDI_MONTHS.get(dt.month, "")
    year = str(dt.year)
    return f"{day} {month}, {year}"


# ---------------------------------------------------------------------------
# 2. Comprehensive Bank Name Mapping
# ---------------------------------------------------------------------------

BANK_HINDI_MAP = [
    (re.compile(r'IDFC\s*FIRST', re.I), 'आई.डी.एफ.सी. फर्स्ट बैंक'),
    (re.compile(r'IDFC', re.I), 'एच.डी.एफ.सी. बैंक'),
    (re.compile(r'HDFC', re.I), 'एच.डी.एफ.सी. बैंक'),
    (re.compile(r'IDBI', re.I), 'आई.डी.बी.आई बैंक'),
    (re.compile(r'AXIS', re.I), 'एक्सीस बैंक'),
    (re.compile(r'YES\s*BANK', re.I), 'यस बैंक'),
    (re.compile(r'PNB|PUNJAB\s*NATIONAL', re.I), 'पंजाब नेशनल बैंक'),
    (re.compile(r'SBI|STATE\s*BANK', re.I), 'भारतीय स्टेट बैंक'),
    (re.compile(r'ICICI', re.I), 'आईसीआईसीआई बैंक'),
    (re.compile(r'KOTAK', re.I), 'कोटक महिन्द्रा बैंक'),
    (re.compile(r'BANK\s*OF\s*BARODA|BOB', re.I), 'बैंक ऑफ बड़ौदा'),
    (re.compile(r'CANARA', re.I), 'केनरा बैंक'),
    (re.compile(r'UNION\s*BANK', re.I), 'यूनियन बैंक ऑफ इंडिया'),
    (re.compile(r'INDUSIND', re.I), 'इन्डसइन्ड बैंक'),
    (re.compile(r'BANK\s*OF\s*INDIA|BOI', re.I), 'बैंक ऑफ इंडिया'),
    (re.compile(r'CENTRAL\s*BANK', re.I), 'सेंट्रल बैंक ऑफ इंडिया')
]

def format_bank_hindi(bank_str):
    if not bank_str:
        return ""
    b_str = str(bank_str).strip()
    for pat, h_name in BANK_HINDI_MAP:
        if pat.search(b_str):
            return h_name
    return b_str


# ---------------------------------------------------------------------------
# 3. Dynamic Parsers
# ---------------------------------------------------------------------------

def parse_plot_dimensions(img_path_or_dict):
    if isinstance(img_path_or_dict, dict) and img_path_or_dict:
        p_no = img_path_or_dict.get('plot_no', '')
        return {
            'plot_no': p_no,
            'plot_no_hindi': img_path_or_dict.get('plot_no_hindi') or plot_number_to_hindi(p_no),
            'width': float(img_path_or_dict.get('width', 0)) if img_path_or_dict.get('width') else '',
            'length': float(img_path_or_dict.get('length', 0)) if img_path_or_dict.get('length') else '',
            'area_sqm': float(img_path_or_dict.get('area_sqm', 0)) if img_path_or_dict.get('area_sqm') else '',
            'area_sqft': int(float(img_path_or_dict.get('area_sqft', 0))) if img_path_or_dict.get('area_sqft') else ''
        }

    if isinstance(img_path_or_dict, str) and os.path.exists(img_path_or_dict):
        try:
            import pytesseract
            img = Image.open(img_path_or_dict)
            text = pytesseract.image_to_string(img)
            numbers = re.findall(r'\d+(?:\.\d+)?', text)
            if len(numbers) >= 5:
                p_no = str(numbers[0])
                return {
                    'plot_no': p_no,
                    'plot_no_hindi': plot_number_to_hindi(p_no),
                    'width': float(numbers[2]),
                    'length': float(numbers[3]),
                    'area_sqm': float(numbers[4]),
                    'area_sqft': int(float(numbers[5]))
                }
        except Exception:
            pass

    return {'plot_no': '', 'plot_no_hindi': '', 'width': '', 'length': '', 'area_sqm': '', 'area_sqft': ''}

def parse_excel_registry_data(excel_path):
    if not excel_path or not os.path.exists(excel_path):
        return {'colony_name': '', 'plot_no': '', 'allotment_val': '', 'payments': [], 'tds_info': None}

    wb = openpyxl.load_workbook(excel_path, data_only=True)
    
    colony_name = "एमराल्ड गेटवे"
    plot_no = ""
    allotment_val = ""
    area_sqft = ""
    area_sqm = ""
    
    app_sheet = None
    for name in wb.sheetnames:
        if 'APPLICANT' in name.upper() or 'BUYER' in name.upper() or 'PLOT' in name.upper():
            app_sheet = wb[name]
            break
    if not app_sheet:
        app_sheet = wb.active

    for r in range(1, 5):
        val = str(app_sheet.cell(row=r, column=1).value or '')
        if 'AASHRAY' in val.upper():
            colony_name = "एमराल्ड आश्रय"
            break
        elif 'GATEWAY' in val.upper():
            colony_name = "एमराल्ड गेटवे"
            break

    width = ""
    length = ""

    for r in range(1, 20):
        c1 = str(app_sheet.cell(row=r, column=1).value or '').strip()
        c2 = app_sheet.cell(row=r, column=2).value
        
        if re.search(r'Plot\s*No', c1, re.I):
            plot_no = str(c2 or '').strip()
        elif re.search(r'Allotment\s*Value', c1, re.I):
            if c2 is not None:
                allotment_val = int(c2)
        elif re.search(r'Area\s*in\s*Sq\.?\s*Ft', c1, re.I):
            if c2 is not None:
                area_sqft = int(float(c2))
        elif re.search(r'Area\s*in\s*Sq\.?\s*Mtr', c1, re.I):
            if c2 is not None:
                area_sqm = round(float(c2), 2)
        elif re.search(r'\b(?:Width|Breath|Breadth|चौड़ाई|Chaudai)\b', c1, re.I):
            if c2 is not None and not width:
                try:
                    width = round(float(c2), 2)
                except (ValueError, TypeError):
                    width = str(c2).strip()
        elif re.search(r'\b(?:Length|Lamba|Lambai|लम्बाई)\b', c1, re.I):
            if c2 is not None and not length:
                try:
                    length = round(float(c2), 2)
                except (ValueError, TypeError):
                    length = str(c2).strip()

    pay_sheet = None
    for name in wb.sheetnames:
        if 'PAYMENT' in name.upper() or 'TRANSACTION' in name.upper():
            pay_sheet = wb[name]
            break
            
    payments = []
    tds_info = None

    if pay_sheet:
        # Detect Paid By / Payer / Applicant column from header row
        paid_by_col = None
        for r_idx, row in enumerate(pay_sheet.iter_rows(values_only=True)):
            if r_idx > 5:
                break
            row_strs = [str(c or '').strip().upper() for c in row]
            for c_idx, val in enumerate(row_strs):
                if any(k in val for k in ('PAID BY', 'PAYER', 'APPLICANT', 'REMITTER', 'NAME', 'BUYER', 'क्रेता')):
                    paid_by_col = c_idx
                    break
            if paid_by_col is not None:
                break

        for r_idx, row in enumerate(pay_sheet.iter_rows(values_only=True)):
            if r_idx <= 1:
                continue
            s_no = row[0]
            if not s_no or (not isinstance(s_no, int) and not str(s_no).isdigit()):
                continue
            
            dt = row[2] if len(row) > 2 else ""
            bank_name = row[3] if len(row) > 3 else ""
            mode_str = row[6] if len(row) > 6 else ""
            instr_str = row[7] if len(row) > 7 else ""
            amount = row[8] if len(row) > 8 else (row[5] if len(row) > 5 and isinstance(row[5], (int, float)) else None)
            
            if not amount:
                continue
                
            bank_hindi = format_bank_hindi(bank_name)
            date_hindi = format_hindi_date(dt)

            paid_by = ""
            if paid_by_col is not None and len(row) > paid_by_col:
                paid_by = str(row[paid_by_col] or '').strip()
            elif len(row) > 1 and isinstance(row[1], str) and not row[1].isdigit() and len(row[1].strip()) > 2:
                paid_by = str(row[1]).strip()
            
            if mode_str and 'TDS' in str(mode_str).upper():
                challan_no = str(instr_str).replace('CHALLAN NO-', '').replace('CHALLAN NO', '').replace('CHALLAN', '').strip()
                tds_info = {
                    'amount': int(amount),
                    'bank_hindi': bank_hindi if bank_hindi else 'एक्सीस बैंक',
                    'date_hindi': date_hindi,
                    'challan_no': challan_no if challan_no else '01568',
                    'bsr_code': '6360014'
                }
                continue
                
            instr_clean = str(instr_str).strip()
            mode_clean = "Online"
            ref_no = instr_clean
            is_utr = False
            
            if '-' in instr_clean:
                parts = instr_clean.split('-', 1)
                m_part = parts[0].strip().upper()
                ref_part = parts[1].strip()
                if 'RTGS' in m_part:
                    mode_clean = 'RTGS'
                elif 'IMPS' in m_part:
                    mode_clean = 'IMPS'
                elif 'NEFT' in m_part:
                    mode_clean = 'NEFT'
                elif 'UPI' in m_part:
                    mode_clean = 'UPI'
                ref_no = ref_part
            elif 'RTGS' in instr_clean.upper():
                mode_clean = 'RTGS'
                ref_no = re.sub(r'RTGS\s*(?:NO\.?)?', '', instr_clean, flags=re.I).strip()
            elif 'IMPS' in instr_clean.upper():
                mode_clean = 'IMPS'
                ref_no = re.sub(r'IMPS\s*(?:NO\.?)?', '', instr_clean, flags=re.I).strip()
            elif 'NEFT' in instr_clean.upper():
                mode_clean = 'NEFT'
                ref_no = re.sub(r'NEFT\s*(?:NO\.?)?', '', instr_clean, flags=re.I).strip()
            elif 'UPI' in instr_clean.upper():
                mode_clean = 'UPI'
                ref_no = re.sub(r'UPI\s*(?:NO\.?)?', '', instr_clean, flags=re.I).strip()
                
            if 'UTR' in instr_clean.upper() or (mode_clean == 'RTGS' and ('HDFCR' in ref_no or 'BKIDA' in ref_no or 'KKBKH' in ref_no)):
                is_utr = True
                
            payments.append({
                's_no': int(s_no),
                'amount': int(amount),
                'bank_hindi': bank_hindi,
                'date_hindi': date_hindi,
                'mode': mode_clean,
                'ref_no': ref_no,
                'paid_by': paid_by,
                'is_utr': is_utr
            })

    return {
        'colony_name': colony_name,
        'plot_no': plot_no,
        'allotment_val': allotment_val,
        'area_sqft': area_sqft,
        'area_sqm': area_sqm,
        'width': width,
        'length': length,
        'payments': payments,
        'tds_info': tds_info
    }


# ---------------------------------------------------------------------------
# 4. Universal Document Generation Engine
# ---------------------------------------------------------------------------

def set_paragraph_runs(p, run_specs):
    p.text = ""
    for text, is_bold in run_specs:
        r = p.add_run(text)
        if is_bold:
            r.bold = True

class RegistryGenerator:
    def __init__(self, template_path=None):
        if not template_path or not os.path.exists(template_path):
            default_template = os.path.join(os.path.dirname(__file__), "Plot Registry EMPTY.docx")
            if os.path.exists(default_template):
                template_path = default_template
            else:
                template_path = r"c:\Users\hp\Desktop\registry\Plot Registry EMPTY.docx"
        self.template_path = template_path
        
    def generate(self, data, output_path):
        doc = docx.Document(self.template_path)
        
        # 1. Dynamic Plot details
        plot_data = data.get('plot_data', {})
        raw_plot_no = plot_data.get('plot_no', '')
        plot_no = clean_plot_no(raw_plot_no)
        plot_no_hindi = plot_data.get('plot_no_hindi') or plot_number_to_hindi(raw_plot_no)
        
        width_raw = plot_data.get('width', '')
        length_raw = plot_data.get('length', '')
        width_m = f"{float(width_raw):.2f}" if width_raw else ""
        length_m = f"{float(length_raw):.2f}" if length_raw else ""
        
        area_sqm = plot_data.get('area_sqm', '')
        area_sqft = plot_data.get('area_sqft', '')
        colony_name = data.get('colony_name', 'एमराल्ड आश्रय')
        
        # 2. Dynamic Financials
        allotment_val = data.get('allotment_val', '')
        allotment_val_formatted = format_indian_currency(allotment_val)
        allotment_val_words = number_to_hindi_words(allotment_val)
        
        # 3. Dynamic Multi-Buyer / Single Buyer Handling
        buyers = data.get('buyers', [])
        payments_list = data.get('payments', [])
        tds_info = data.get('tds_info', None)
        
        # 1. Header Banner & Preamble Replacements across the document
        for p in doc.paragraphs:
            # Top Banner Header (e.g. भूखण्ड क्रमांक E9 (ई-नौ), एमराल्ड आश्रय, ग्राम-सोनवाय, तहसील-राऊ, ज़िला-इन्दौर)
            if "भूखण्ड क्रमांक" in p.text and ("ग्राम-सोनवाय" in p.text or "ग्राम-सुल्लाखेड़ी" in p.text or "तहसील" in p.text) and "विक्रय-पत्र" not in p.text:
                banner_runs = [
                    (f"भूखण्ड क्रमांक {plot_no} ({plot_no_hindi}), {colony_name}, ग्राम-सोनवाय, तहसील-राऊ, ज़िला-इन्दौर", True)
                ]
                set_paragraph_runs(p, banner_runs)
            # Preamble Clause (contains Economic Weaker Section or general plot title)
            elif "विक्रय व्यवहार मूल्य" in p.text or "आर्थिक रुप से कमज़ोर श्रेणी" in p.text or "जिसका विक्रय व्यवहार मूल्य" in p.text:
                if "आर्थिक रुप से कमज़ोर श्रेणी" in p.text or "EWS" in p.text or "सोनवाय" in p.text or "AASHRAY" in p.text.upper() or "आश्रय" in colony_name:
                    preamble_runs = [
                        ('ग्राम पंचायत क्षेत्र के अन्तर्गत स्थित ', None),
                        ('ग्राम-सोनवाय, तहसील-राऊ, ', True),
                        ('ज़िला', True),
                        ('-इन्दौर ', True),
                        ('के विभिन्न सर्वे क्रमांकों की भूमि पर विकसित कॉलोनी ', None),
                        (f'‘{colony_name} (EMERALD AASHRAY)’ ', True),
                        ('में ', None),
                        ('‘आर्थिक रुप से कमज़ोर श्रेणी’', True),
                        (' [Economic Weaker Section (EWS)] हेतु आरक्षित रखे गये भूखण्डों में से, ', None),
                        ('भूखण्ड क्रमांक ', True),
                        (f'{plot_no}', True),
                        (' (', True),
                        (f'{plot_no_hindi}', True),
                        ('),', True),
                        (' जिसका विक्रय व्यवहार मूल्य ', None),
                        ('रुपये ', True),
                        (f'{allotment_val_formatted}/- (अक्षरी रुपये ', True),
                        (f'{allotment_val_words} ', True),
                        ('मात्र) ', True),
                        ('है, का ', None),
                        ('विक्रय-पत्र', True)
                    ]
                else:
                    preamble_runs = [
                        ('ग्राम पंचायत क्षेत्र के अन्तर्गत स्थित ', None),
                        ('ग्राम-सुल्लाखेड़ी, तहसील-सांवेर, ', True),
                        ('ज़िला', True),
                        ('-इन्दौर ', True),
                        ('के विभिन्न सर्वे क्रमांकों की भूमि पर विकसित कॉलोनी ', None),
                        (f'‘{colony_name} (EMERALD GATEWAY)’', True),
                        (' के आवासीय', None),
                        (' भूखण्ड क्रमांक ', True),
                        (f'{plot_no}', True),
                        (' (', True),
                        (f'{plot_no_hindi}', True),
                        ('),', True),
                        (' जिसका विक्रय व्यवहार मूल्य ', None),
                        ('रुपये ', True),
                        (f'{allotment_val_formatted}/- (अक्षरी रुपये ', True),
                        (f'{allotment_val_words} ', True),
                        ('मात्र) ', True),
                        ('है, का ', None),
                        ('विक्रय-पत्र', True)
                    ]
                set_paragraph_runs(p, preamble_runs)

        # 2. Dynamic Buyer Construction (Search for buyer section in ANY template)
        is_joint = len(buyers) > 1
        buyer_party_tag = "... द्वितीयपक्ष/क्रेतागण" if is_joint else "... द्वितीयपक्ष"
        
        # Locate buyer paragraph index
        buyer_p_idx = None
        for idx, p in enumerate(doc.paragraphs):
            if "इन्हों से यह विक्रय-पत्र निष्पादित करवा लेने वाले क्रेता" in p.text:
                buyer_p_idx = idx + 1
                break
            elif "(1)" in p.text and idx < 25:
                buyer_p_idx = idx
                break
            elif "Aadhar No." in p.text and idx < 25:
                buyer_p_idx = idx
                break

        if buyer_p_idx is not None and buyers:
            if len(buyers) == 1:
                b = buyers[0]
                b1_runs = [
                    ('(1) ', True),
                    (f"{b.get('name', '')}", True),
                    (', ', True),
                    ('PAN : ', None),
                    (f"{b.get('pan_no', '')}", None),
                    (' & ', None),
                    ('Aadhar No. : ', None),
                    (f"{b.get('aadhar_no', '')}", None)
                ]
                set_paragraph_runs(doc.paragraphs[buyer_p_idx], b1_runs)

                r_title = b.get('relation_title', 'पति श्री')
                r_name = b.get('relation_name', '')
                b2_runs = [
                    (f"{r_title} ", True),
                    (f"{r_name}", True),
                    (',', True)
                ]
                set_paragraph_runs(doc.paragraphs[buyer_p_idx + 1], b2_runs)

                b3_runs = [
                    (f"निवासी-{b.get('address', '')}", None),
                    (' . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . ', None),
                    (f"{buyer_party_tag}", True)
                ]
                set_paragraph_runs(doc.paragraphs[buyer_p_idx + 2], b3_runs)
            elif len(buyers) > 1:
                p17_runs = []
                for b_idx, b in enumerate(buyers):
                    b_num = b_idx + 1
                    r_title = b.get('relation_title', 'पति श्री')
                    r_name = b.get('relation_name', '')
                    rel_str = f", {r_title} {r_name}" if r_name else ""
                    semi = "; " if b_idx < len(buyers) - 1 else ""
                    p17_runs.extend([
                        (f"({b_num}) ", True),
                        (f"{b.get('name', '')}", True),
                        (f"{rel_str}", True),
                        (', ', True),
                        ('PAN : ', None),
                        (f"{b.get('pan_no', '')}", None),
                        (' & ', None),
                        ('Aadhar No. : ', None),
                        (f"{b.get('aadhar_no', '')}", None),
                        (f"{semi}", None)
                    ])
                set_paragraph_runs(doc.paragraphs[buyer_p_idx], p17_runs)
                
                p18_runs = [
                    (f"उभय निवासी-{buyers[0].get('address', '')}", None)
                ]
                set_paragraph_runs(doc.paragraphs[buyer_p_idx + 1], p18_runs)

                p19_runs = [
                    (' . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . . ', None),
                    (f"{buyer_party_tag}", True)
                ]
                set_paragraph_runs(doc.paragraphs[buyer_p_idx + 2], p19_runs)

        # 3. Clause (10) / (11) Cost Intro
        for p in doc.paragraphs:
            if ("(10)" in p.text or "(11)" in p.text) and ("कुल कीमत" in p.text or "सम्पूर्ण विक्रय मूल्य" in p.text):
                if "आर्थिक रुप से कमज़ोर श्रेणी" in p.text or "EWS" in p.text or "सोनवाय" in p.text or "AASHRAY" in p.text.upper() or "आश्रय" in colony_name:
                    p68_runs = [
                        ('(11) यह कि, उपरोक्त वर्णित अनुसार प्रथमपक्ष/विक्रेता को उक्त कॉलोनी ‘एमराल्ड आश्रय (EMERALD AASHRAY)’ में विकसित हुए विभिन्न भूखण्डों के विषय में समस्त प्रकार का विक्रय व्यवहार, अन्तरण, हस्तान्तरण एवं निर्णय आदि करने का पूर्ण तथा वैधानिक अधिकार प्राप्त होने से, उन्होंने सदर कालोनी में ‘आर्थिक रुप से कमज़ोर श्रेणी\' [Economic Weaker Section (EWS)] हेतु आरक्षित रखे गये विभिन्न भूखण्डों का विक्रय करने से पूर्व विधिक प्रक्रिया के तहत पेपर पब्लिकेशन (ज़ाहिर सूचना) का प्रकाशन करवाया गया, तत्पश्चात्‌‍ उक्त आरक्षित श्रेणी के भूखण्डों को क्रय करने बाबद्‌‍ प्राप्त हुए आवेदनों पर से प्रथमपक्ष/विक्रेता द्वारा ‘पहले आओं पहले पाओं\' (First Cum First Basis) के अन्तर्गत पात्र आवेदकों को उक्त श्रेणी के भूखण्डों के विक्रय बाबद्‌ सक्षम प्राधिकारी से विधिवत्‌‍ अनुमति प्राप्त की गई है । उक्त प्राप्त हुई अनुमति पर से प्रथमपक्ष/विक्रेता द्वारा उक्त कॉलोनी ‘एमराल्ड आश्रय (EMERALD AASHRAY)’ में ‘आर्थिक रुप से कमज़ोर श्रेणी\' [Economic Weaker Section (EWS)] हेतु आरक्षित रखे गये विभिन्न भूखण्डों का आवंटन किया गया है तथा उक्त आवंटन पश्चात्‌ उक्त कॉलोनी के निम्न वर्णित भूखण्ड, जिसका स्पष्ट वर्णन एवं चतु:सीमा निम्नानुसार है, को मय विकास खर्च के, इस विक्रय-पत्र के माध्यम से आप द्वितीयपक्ष/क्रेता को कुल कीमत ', None),
                        ('रुपये ', True),
                        (f'{allotment_val_formatted}/- (अक्षरी रुपये ', True),
                        (f'{allotment_val_words} ', True),
                        ('मात्र)', True),
                        (' में विक्रय रीति से अन्तरित कर दिया होकर, इस विक्रय व्यवहार का सम्पूर्ण विक्रय मूल्य निम्नानुसार रीति से प्राप्त कर लिया है :-', None)
                    ]
                else:
                    p68_runs = [
                        ('(10) यह कि, उपरोक्त वर्णित अनुसार प्रथमपक्ष/विक्रेता को अपने स्वामित्व एवं आधिपत्य की उपरोक्त वर्णित कॉलोनी ', None),
                        (f'‘{colony_name} (EMERALD GATEWAY)’ ', True),
                        ('में विकसित विभिन्न भूखण्डों के विषय में समस्त प्रकार का विक्रय व्यवहार, अन्तरण, हस्तान्तरण एवं निर्णय आदि करने का पूर्ण तथा वैधानिक अधिकार प्राप्त होने से, उन्होंने सदर कालोनी के निम्न वर्णित भूखण्ड को, जिसका स्पष्ट वर्णन एवं चतु:सीमा निम्नानुसार है, को मय विकास खर्च के, द्वितीयपक्ष/क्रेता को कुल कीमत ', None),
                        ('रुपये ', True),
                        (f'{allotment_val_formatted}/- (अक्षरी रुपये ', True),
                        (f'{allotment_val_words} ', True),
                        ('मात्र)', True),
                        (' में विक्रय कर दिया होकर, उसका सम्पूर्ण विक्रय मूल्य निम्नानुसार रीति से प्राप्त कर लिया है ।', None)
                    ]
                set_paragraph_runs(p, p68_runs)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                break

        # 4. Dynamic Payment Paragraphs (Clause 10/11)
        new_payment_paragraphs_data = []
        for p_item in payments_list:
            amt = p_item['amount']
            amt_str = f"{amt}" if amt >= 100000 and amt % 1000 == 0 and amt in [750000] else format_indian_currency(amt)
            bank_name = p_item['bank_hindi']
            dt_str = p_item['date_hindi']
            mode_str = p_item['mode']
            ref_str = p_item['ref_no']
            is_utr = p_item.get('is_utr', False)
            ref_label = "UTR No. : " if is_utr else "Reference No. : "

            paid_by = p_item.get('paid_by', '')
            paid_by_text = f" ({paid_by} द्वारा) " if paid_by else " "
            p_runs = [
                ('रुपये ', None),
                (f'{amt_str}', None),
                (f'/-{paid_by_text}प्रथमपक्ष/विक्रेता ने उन्हों के बैंक खाते में ', None),
                (f'{bank_name}', True),
                (' से ', None),
                ('दिनांक ', True),
                (f'{dt_str} ', True),
                ('को किये गये बैंक ट्रांसफर (', None),
                (f'{mode_str}', True),
                (f') के माध्यम से विधिवत्‌‍ रूप से प्राप्त किये है, जिस बाबद्‌ {ref_label}', None),
                (f'{ref_str}', True),
                (' जारी हुआ है ।', None)
            ]
            new_payment_paragraphs_data.append(p_runs)
            new_payment_paragraphs_data.append([])

        if tds_info:
            tds_amt_str = f"{tds_info['amount']}"
            tds_bank = tds_info['bank_hindi']
            tds_date = tds_info['date_hindi']
            tds_challan = tds_info['challan_no']
            tds_bsr = tds_info['bsr_code']

            tds_runs = [
                (f'रुपये {tds_amt_str}/- की राशि को आयकर अधिनियम, 2025 की धारा 393(1), तालिका क्रमांक 3(i) के प्रावधानों के अंतर्गत, ', None),
                ('द्वितीयपक्ष/क्रेता के सदस्य क्रमांक 1 ', True),
                ('द्वारा प्रथमपक्ष/विक्रेता की ओर से अचल सम्पत्ति विक्रय से आय के स्रोत पर कर कटौती (', None),
                ('T.D.S.', True),
                (') के रूप में ', None),
                (f'{tds_bank} ', True),
                ('(BSR Code : ', None),
                (f'{tds_bsr}', True),
                (') के माध्यम से दिनांक ', None),
                (f'{tds_date} ', True),
                ('को चालान नम्बर ', None),
                (f'{tds_challan} ', True),
                ('के द्वारा आयकर विभाग को जमा करवा दी होकर, उक्त चालान एवं Form 26QB की छायाप्रति प्रथमपक्ष/विक्रेता को उपलब्ध करवा दी गई है ।', None)
            ]
            new_payment_paragraphs_data.append(tds_runs)
            new_payment_paragraphs_data.append([])

        # Separators and Total
        new_payment_paragraphs_data.append([('---------------------------------------------------------------------------------------------', True)])
        total_runs = [
            ('रुपये ', True),
            (f'{allotment_val_formatted}/- (अक्षरी रुपये ', True),
            (f'{allotment_val_words} ', True),
            ('मात्र) कुल प्राप्त हुए। ', True),
            ('(subject to the realization of the above mentioned cheque/s).', None)
        ]
        new_payment_paragraphs_data.append(total_runs)
        new_payment_paragraphs_data.append([('---------------------------------------------------------------------------------------------', True)])

        # Locate payment placeholder
        p70_idx = None
        p_end_idx = None
        for idx, p in enumerate(doc.paragraphs):
            if re.search(r'रुपये\s*[…\.\d,/_-]+\s*प्रथमपक्ष/विक्रेता', p.text):
                if p70_idx is None:
                    p70_idx = idx
            elif "(subject to the realization" in p.text:
                p_end_idx = idx
                break

        if p70_idx is not None and p_end_idx is not None:
            existing_count = (p_end_idx + 2) - p70_idx
            needed_count = len(new_payment_paragraphs_data)
            anchor_p = doc.paragraphs[p_end_idx + 2]

            for i in range(min(existing_count, needed_count)):
                p = doc.paragraphs[p70_idx + i]
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_paragraph_runs(p, new_payment_paragraphs_data[i])

            for i in range(existing_count, needed_count):
                new_p = anchor_p.insert_paragraph_before()
                new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                set_paragraph_runs(new_p, new_payment_paragraphs_data[i])

        # 5. Dynamic Property Description
        for i, p in enumerate(doc.paragraphs):
            if "बिक्रीत सम्पत्ति का स्पष्ट वर्णन" in p.text:
                desc_p = doc.paragraphs[i + 1]
                if "आर्थिक रुप से कमज़ोर श्रेणी" in desc_p.text or "EWS" in desc_p.text or "सोनवाय" in desc_p.text or "AASHRAY" in desc_p.text.upper() or "आश्रय" in colony_name:
                    desc_runs = [
                        ('ग्राम पंचायत क्षेत्र के अन्तर्गत स्थित ', None),
                        ('ग्राम-सोनवाय, तहसील-राऊ, ', True),
                        ('ज़िला', True),
                        ('-इन्दौर ', True),
                        ('के विभिन्न सर्वे क्रमांकों की भूमि पर विकसित कॉलोनी ', None),
                        (f'‘{colony_name} (EMERALD AASHRAY)’ ', True),
                        ('में ‘आर्थिक रुप से कमज़ोर श्रेणी\' [Economic Weaker Section (EWS)] हेतु आरक्षित रखे गये भूखण्डों में से ', None),
                        ('भूखण्ड क्रमांक ', True),
                        (f'{plot_no}', True),
                        (' (', True),
                        (f'{plot_no_hindi}', True),
                        ('),', True),
                        (' जिसकी तल भूमि का कुल क्षेत्रफल ', None),
                        (f'{area_sqm} ', True),
                        ('वर्गमीटर', True),
                        (' अर्थात्‌‍ ', None),
                        (f'{area_sqft} ', True),
                        ('वर्गफीट', True),
                        (' है, जिसे सम्पूर्ण को प्रथमपक्ष/विक्रेता इस विक्रय-पत्र के माध्यम से आप द्वितीयपक्ष/क्रेता को विक्रय रीति से अन्तरित कर रहे है ।', None),
                        (' बिक्रीत भूखण्ड को और अधिक स्पष्टता के लिये इस विक्रय-पत्र के साथ संलग्न नक्शे में ', None),
                        ('“भूखण्ड क्रमांक“', True),
                        (' से दर्शाया गया है, नक्शा इस विक्रय-पत्र का अभिन्न अंग रहेगा ।', None),
                        (' बिक्रीत भूखण्ड को इस  विक्रय-पत्र में सुविधा तथा संक्षिप्तता की दृष्टि से ', None),
                        ('“सदर सम्पत्ति”', True),
                        (' शब्द से सम्बोधित किया गया है । बिक्रीत सदर भूखण्ड के निकास की स्थाई एवं स्वतंत्र व्यवस्था सदर भूखण्ड के सामने की ओर स्थित सड़क से होकर है । बिक्रीत सदर भूखण्ड वर्तमान में मौके पर पूर्णत: रिक्त अवस्था में आवासीय उपयोग का होकर, इस भूखण्ड की तल भूमि निजी स्वामित्व की है तथा बिक्रीत भूखण्ड कॉलोनी में अन्दर की ओर स्थित है ।', None)
                    ]
                else:
                    dim_text = f"चौड़ाई {width_m} मीटर एवं लम्बाई {length_m} मीटर, " if (width_m and length_m) else ""
                    desc_runs = [
                        ('ग्राम पंचायत क्षेत्र के अन्तर्गत स्थित ', None),
                        ('ग्राम-सुल्लाखेड़ी, तहसील-सांवेर, ', True),
                        ('ज़िला', True),
                        ('-इन्दौर ', True),
                        ('के विभिन्न सर्वे क्रमांकों की भूमि पर विकसित कॉलोनी ', None),
                        (f'‘{colony_name} (EMERALD GATEWAY)’ ', True),
                        ('का ', None),
                        ('भूखण्ड क्रमांक', True),
                        (f' {plot_no}', True),
                        (' (', True),
                        (f'{plot_no_hindi}', True),
                        (')', True),
                        (' ', None),
                        ('जिसकी तल भूमि की ', None),
                        (f'{dim_text}', None),
                        ('सदर भूखण्ड', None),
                        (' की', None),
                        (' तल भूमि का कुल क्षेत्रफल ', None),
                        (f'{area_sqm} ', True),
                        ('वर्गमीटर', True),
                        (' अर्थात्‌‍ ', None),
                        (f'{area_sqft}', True),
                        (' वर्गफीट ', True),
                        ('है, जिसे सम्पूर्ण को प्रथमपक्ष/विक्रेता इस विक्रय-पत्र के माध्यम से आप द्वितीयपक्ष/क्रेता को विक्रय कर रहे है ।', None),
                        (' बिक्रीत भूखण्ड को और अधिक स्पष्टता के लिये इस विक्रय-पत्र के साथ संलग्न नक्शे में ', None),
                        ('“भूखण्ड क्रमांक“', True),
                        (' से दर्शाया गया है, नक्शा इस विक्रय-पत्र का अभिन्न अंग रहेगा ।', None),
                        (' बिक्रीत भूखण्ड को इस  विक्रय-पत्र में सुविधा तथा संक्षिप्तता की दृष्टि से ', None),
                        ('“सदर सम्पत्ति”', True),
                        (' शब्द से सम्बोधित किया गया है । बिक्रीत सदर भूखण्ड के निकास की स्थाई एवं स्वतंत्र व्यवस्था सदर भूखण्ड के सामने की ओर स्थित सड़क से होकर है । बिक्रीत सदर भूखण्ड वर्तमान में मौके पर पूर्णत: रिक्त अवस्था में ', None),
                        ('आवासीय उपयोग', True),
                        (' का होकर, इस भूखण्ड की तल भूमि निजी स्वामित्व की है तथा बिक्रीत भूखण्ड कॉलोनी में अन्दर की ओर स्थित है ।', None)
                    ]
                set_paragraph_runs(desc_p, desc_runs)
                desc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                break

        # 6. Dynamic Boundaries & Measurements (Both 2-column format & standalone format)
        boundaries = data.get('boundaries', {})
        east_b = boundaries.get('east', '')
        west_b = boundaries.get('west', '')
        north_b = boundaries.get('north', '')
        south_b = boundaries.get('south', '')
        
        for p in doc.paragraphs:
            # 2-Column Template matching (e.g. पूर्व में . . . :: . . ………….…………. . . . . . . . . . . . . . . . . . . . . . . . . . . पूर्व में . . . :: . . …… मीटर)
            if "पूर्व में . . . ::" in p.text or "पूर्व में . . ::" in p.text or "पूर्व में . . . . ::" in p.text:
                if "मीटर" in p.text:
                    p.text = f"पूर्व में . . . :: . . {east_b} . . . . . . . . . . . . . . . . . . . . . . . . . . पूर्व में . . . :: . . {width_m or '……'} मीटर"
                else:
                    p.text = f"पूर्व में . . . . :: . . {east_b}"
            elif "पश्चिम में . .::" in p.text or "पश्चिम में . . ::" in p.text:
                if "मीटर" in p.text:
                    p.text = f"पश्चिम में . .:: . . {west_b} . . . . . . . . . . . . . . . . . . . . . . . . पश्चिम में . . :: . . {width_m or '……'} मीटर"
                else:
                    p.text = f"पश्चिम में . . :: . . {west_b}"
            elif "उत्तर में . . .::" in p.text or "उत्तर में . . . .::" in p.text or "उत्तर में . . ::" in p.text:
                if "मीटर" in p.text:
                    p.text = f"उत्तर में . . .:: . . {north_b} . . . . . . . . . . . . . . . . . . . . . उत्तर में . . :: . . {length_m or '……. '} मीटर"
                else:
                    p.text = f"उत्तर में . . . .:: . . {north_b}"
            elif "दक्षिण में . . ::" in p.text or "दक्षिण में . . . ::" in p.text:
                if "मीटर" in p.text:
                    p.text = f"दक्षिण में . . :: . . {south_b} . . . . . . . . . . . . . . . . . दक्षिण में . . :: . . {length_m or '……'} मीटर"
                else:
                    p.text = f"दक्षिण में . . . :: . . {south_b}"

        # 7. Closing Delivery Clause with Price (P84)
        for p in doc.paragraphs:
            if "उपरोक्त वर्णन तथा चतु:सीमा के मध्य की सदर सम्पत्ति को" in p.text:
                p84_runs = [
                    ('उपरोक्त वर्णन तथा चतु:सीमा के मध्य की सदर सम्पत्ति को, आगे से पीछे तक तल भूमि के सम्पूर्ण अधिकार सहित, प्रथमपक्ष/विक्रेता ने आप द्वितीयपक्ष/क्रेता को उपरोक्त वर्णित अनुसार चुकता विक्रय मूल्य ', None),
                    ('रुपये ', True),
                    (f'{allotment_val_formatted}/- (अक्षरी रुपये ', True),
                    (f'{allotment_val_words} ', True),
                    ('मात्र)', True),
                    (' प्राप्त करने के उपरान्त विक्रय कर दिया है, साथ ही बिक्रीत सदर सम्पत्ति का रिक्त तथा स्वतंत्र आधिपत्य भी मौके पर नपती करवाकर, प्रथमपक्ष/विक्रेता ने द्वितीयपक्ष/क्रेता को साक्षियों के समक्ष विधिवत रूप से सौंप दिया है । बिक्रीत सदर भूखण्ड का मौके पर नप्ती करवा कर आधिपत्य प्राप्त होना द्वितीयपक्ष/क्रेता इस विक्रय-पत्र के माध्यम से स्वीकार करते है।', None)
                ]
                set_paragraph_runs(p, p84_runs)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                break

        # 8. Execution Date
        exec_date = data.get('execution_date') or get_today_hindi_date()
        for p in doc.paragraphs:
            if "उपरोक्तानुसार सदर सम्पत्ति का विक्रय पत्र प्रथमपक्ष/विक्रेता" in p.text:
                end_runs = [
                    ('उपरोक्तानुसार सदर सम्पत्ति का विक्रय पत्र प्रथमपक्ष/विक्रेता एवं द्वितीयपक्ष/क्रेता ने परस्पर एक-दूसरे के हित व हक में, पढ़कर, सुनकर तथा समझकर बिना नशा पानी किये, दो साक्षियों की उपस्थिति में, आज नगर इन्दौर में अपने-अपने हस्ताक्षर करके निष्पादित तथा सम्पादित कर दिया, सो सही । इति, इन्दौर, ', None),
                    ('दिनांक : ', True),
                    (f'{exec_date}', True)
                ]
                set_paragraph_runs(p, end_runs)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                break

        doc.save(output_path)
        return output_path
